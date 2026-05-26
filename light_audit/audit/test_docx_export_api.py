"""Tests for US-051: Narrative .docx export."""
import io
from http import HTTPStatus

import docx
import pytest
from django.contrib.auth import get_user_model

from light_audit.audit.models import AgentRun
from light_audit.audit.models import AgentType
from light_audit.audit.models import AuditFlag
from light_audit.audit.models import AuditVersion
from light_audit.audit.models import Building
from light_audit.audit.models import Floor
from light_audit.audit.models import LogEntry
from light_audit.audit.models import Project
from light_audit.audit.models import Room

User = get_user_model()


@pytest.fixture
def user(db):
    return User.objects.create_user(email="docx@example.com", password="pass")  # noqa: S106


@pytest.fixture
def project(db, user):
    return Project.objects.create(name="Docx Project", owner=user)


@pytest.fixture
def building(db, project):
    return Building.objects.create(name="Docx Building", project=project)


@pytest.fixture
def version(db, building, user):
    return AuditVersion.objects.create(building=building, created_by=user)


@pytest.fixture
def floor(db, building, version):
    return Floor.objects.create(
        building=building, audit_version=version, name="Ground", level=1,
    )


@pytest.fixture
def room(db, floor, version):
    return Room.objects.create(
        floor=floor, audit_version=version, name="Office 101",
    )


@pytest.fixture
def log_entry(db, room, version):
    return LogEntry.objects.create(
        room=room, audit_version=version, fixture_id="F-001",
    )


@pytest.fixture
def agent_run(db, version, user, project):
    run = AgentRun.objects.create(
        agent_type=AgentType.AUDIT_REVIEW,
        user=user,
        project=project,
        audit_version=version,
        status="ok",
    )
    flags_json = (
        '[{"log_entry_id": 1, "severity": "high", "message": "replace fixture"}]'
    )
    run.response_output = {
        "content": (
            "This building has several issues.\n\n"
            "Energy efficiency is low.\n"
            f"```flags\n{flags_json}\n```"
        ),
    }
    run.save()
    return run


@pytest.fixture
def api_client(client, user):
    client.force_login(user)
    return client


def test_docx_auth_required(client, version):
    resp = client.post(f"/api/audit-versions/{version.pk}/export/docx/")
    assert resp.status_code == HTTPStatus.UNAUTHORIZED


def test_docx_returns_file(api_client, version, agent_run):
    resp = api_client.post(f"/api/audit-versions/{version.pk}/export/docx/")
    assert resp.status_code == HTTPStatus.OK
    assert "wordprocessingml" in resp["Content-Type"]
    assert f"audit-narrative-{version.pk}.docx" in resp["Content-Disposition"]


def test_docx_404_no_run(api_client, version):
    """No AgentRun → 404."""
    resp = api_client.post(f"/api/audit-versions/{version.pk}/export/docx/")
    assert resp.status_code == HTTPStatus.NOT_FOUND


def test_docx_404_unknown_version(api_client):
    resp = api_client.post("/api/audit-versions/99999/export/docx/")
    assert resp.status_code == HTTPStatus.NOT_FOUND


def test_docx_contains_summary(api_client, version, agent_run):
    resp = api_client.post(f"/api/audit-versions/{version.pk}/export/docx/")
    doc = docx.Document(io.BytesIO(resp.content))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "This building has several issues" in text
    assert "Energy efficiency is low" in text


def test_docx_flags_block_stripped_from_summary(api_client, version, agent_run):
    """The raw ```flags``` block must NOT appear in the docx text."""
    resp = api_client.post(f"/api/audit-versions/{version.pk}/export/docx/")
    doc = docx.Document(io.BytesIO(resp.content))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "```flags" not in text


def test_docx_includes_flag_details(api_client, version, agent_run, log_entry):
    AuditFlag.objects.create(
        log_entry=log_entry,
        audit_version=version,
        severity="high",
        message="Replace this fixture immediately.",
        source_run=agent_run,
    )
    resp = api_client.post(f"/api/audit-versions/{version.pk}/export/docx/")
    doc = docx.Document(io.BytesIO(resp.content))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Replace this fixture immediately." in text
    assert "HIGH" in text


def test_docx_dismissed_flags_excluded(api_client, version, agent_run, log_entry, user):
    flag = AuditFlag.objects.create(
        log_entry=log_entry,
        audit_version=version,
        severity="low",
        message="This should be hidden.",
        source_run=agent_run,
    )
    flag.dismiss(user=user, reason="not relevant")
    resp = api_client.post(f"/api/audit-versions/{version.pk}/export/docx/")
    doc = docx.Document(io.BytesIO(resp.content))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "This should be hidden." not in text


def test_docx_no_flags_message(api_client, version, agent_run):
    """When no active flags, document says so."""
    resp = api_client.post(f"/api/audit-versions/{version.pk}/export/docx/")
    doc = docx.Document(io.BytesIO(resp.content))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "No active flags" in text
